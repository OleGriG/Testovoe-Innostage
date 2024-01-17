import docx
import re
import string
from memory_profiler import profile, memory_usage
import timeit


def count_words(file_name):
    document = docx.Document(file_name)
    paragraphs_list = []
    for paragraph_text in document.paragraphs:
        paragraphs_list.append(paragraph_text.text)

    return len(re.findall(r'\w+', '\n'.join(paragraphs_list)))


def count_symbols(file_path):
    punctuation_marks = set(string.punctuation)
    doc = docx.Document(file_path)
    symbol_count = 0
    for paragraph in doc.paragraphs:
        cleaned_text = ''.join(char for char in paragraph.text if char != '\b' and char not in punctuation_marks and char != ' ')
        symbol_count += len(cleaned_text)
    return symbol_count


@profile
def fibonacci_list(n):
    fib_list = [0, 1]
    while len(fib_list) < n:
        fib_list.append(fib_list[-1] + fib_list[-2])
    return fib_list[:n]


def fibonacci_generator(n):
    a, b = 0, 1
    count = 0
    while count < n:
        yield a
        a, b = b, a + b
        count += 1


if __name__ == '__main__':
    print(count_words("text_task.docx"))
    print(count_symbols("text_task.docx"))
    # Я не нашел никакого варианта посчитать именно строки в конкретном представлении ворда,
    # потому что учитывается размер символов, настройки документа с полями и прочие условия

    time_list = timeit.timeit(lambda: fibonacci_list(100), number=1)
    print(f"Время выполнения обычной функции: {time_list} секунд")
    print(f'память до: {memory_usage()}MB')
    time_gen = timeit.timeit(lambda: list(fibonacci_generator(100)), number=1) 
    print(f'память после: {memory_usage()}MB')
    # По моему задание некорректно, т.к. генератор в принципе, ничего не
    # хранит в памяти кроме текущего состояния, в том и его прелесть,
    # чтоб не быть голословным добавил вывод использования памяти до и после
    print(f"Время выполнения функции-генератора: {time_gen} секунд")


# Доп задание:
# ошибка рейзится внтури with, а метод __exit__ не обрабатывает ее
# __exit__ должен вернуть True, если в блоке with исключения были обработаны внутри __exit__
# Если в блоке with возникло исключение, которое не обрабатывается в __exit__,
# или не было исключений вообще, то функция должна возвращать False
